"""
DOCX document builder per documentation.mdc best practices.
Converts Markdown to DOCX format with proper formatting.
"""
import io
import json
import re
import logging
from typing import List

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.table import Table

logger = logging.getLogger(__name__)

def build_docx_document(markdown_text: str) -> bytes:
    """
    Сконвертировать готовый Markdown в документ DOCX.

    Args:
        markdown_text: Markdown-представление документации.

    Returns:
        bytes: Содержимое DOCX-файла.
    """
    document = Document()
    ensure_code_style(document)
    lines = markdown_text.splitlines()
    index = 0
    previous_blank = True
    previous_was_parameter_header = False  # Отслеживание, была ли предыдущая строка "Параметры:", "Возвращает:" или "Вызывает:"

    while index < len(lines):
        line = lines[index]
        stripped = line.strip()

        if not stripped:
            if not previous_blank:
                document.add_paragraph("")
                previous_blank = True
            # Не сбрасываем флаг, если следующая строка может быть таблицей
            # Флаг сбросится при обработке таблицы или другого элемента
            index += 1
            continue

        if stripped.startswith("#"):
            heading_level = min(len(stripped) - len(stripped.lstrip("#")), 5)
            heading_text = stripped.lstrip("#").strip()
            document.add_heading(heading_text or line, level=heading_level)
            previous_blank = False
            previous_was_parameter_header = False
            index += 1
            continue

        if stripped.startswith("|") and stripped.endswith("|"):
            table_block: List[str] = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_block.append(lines[index])
                index += 1
            # Если таблица идет после заголовков Parameters/Returns/Raises, создаем без рамок
            add_table_from_markdown(document, table_block, no_borders=previous_was_parameter_header)
            previous_blank = False
            previous_was_parameter_header = False
            continue

        if stripped.startswith("```"):
            code_lines: List[str] = []
            index += 1
            while index < len(lines) and not lines[index].strip().startswith("```"):
                code_lines.append(lines[index])
                index += 1
            code_content = "\n".join(code_lines).strip()
            try:
                formatted_json = json.dumps(json.loads(code_content), ensure_ascii=False, indent=2)
            except json.JSONDecodeError:
                formatted_json = code_content
            add_code_block(document, formatted_json)
            index += 1  # skip closing ```
            previous_blank = False
            previous_was_parameter_header = False
            continue

        if stripped == "---":
            if not previous_blank:
                document.add_paragraph("")
                previous_blank = True
            previous_was_parameter_header = False
            index += 1
            continue

        if stripped.startswith("- ") or stripped.startswith("* "):
            add_list_item(document, line, stripped, reduce_indent=previous_was_parameter_header)
            previous_blank = False
            previous_was_parameter_header = False
            index += 1
            continue

        # Специальная обработка для строк "Параметры:", "Возвращает:", "Вызывает:" 
        # чтобы убрать большие отступы перед списками
        if stripped.endswith(":") and any(keyword.lower() in stripped.lower() for keyword in ["Параметры", "Возвращает", "Вызывает", "Parameters", "Returns", "Raises"]):
            # Переводим английские заголовки на русский
            translated_line = line
            # Проверяем и переводим каждый возможный английский заголовок
            for eng_base, rus_base in [("Parameters", "Параметры"), ("Returns", "Возвращает"), ("Raises", "Вызывает")]:
                # Ищем английский заголовок в строке (с двоеточием) - более точный паттерн
                # Проверяем, что строка начинается с заголовка или содержит его как отдельное слово
                pattern = re.compile(rf'^\s*{re.escape(eng_base)}\s*:|\b{re.escape(eng_base)}\s*:', re.IGNORECASE)
                if pattern.search(stripped):
                    # Проверяем, что русский вариант еще не присутствует
                    if rus_base not in stripped:
                        # Заменяем английский заголовок на русский, сохраняя пробелы в начале
                        leading_spaces = len(line) - len(line.lstrip())
                        translated_line = " " * leading_spaces + pattern.sub(f"{rus_base}:", stripped.lstrip())
                        break
            
            paragraph = document.add_paragraph(translated_line)
            p_pr = paragraph._element.get_or_add_pPr()
            
            # Устанавливаем spacing - для Returns и Raises добавляем отступ сверху
            existing_spacing = p_pr.find(qn("w:spacing"))
            if existing_spacing is not None:
                p_pr.remove(existing_spacing)
            spacing = OxmlElement("w:spacing")
            # Для Returns и Raises добавляем отступ сверху
            if any(keyword in stripped for keyword in ["Возвращает", "Вызывает", "Returns", "Raises"]):
                spacing.set(qn("w:before"), "240")  # Отступ перед заголовком
            else:
                spacing.set(qn("w:before"), "0")
            spacing.set(qn("w:after"), "0")
            spacing.set(qn("w:line"), "240")  # Минимальный межстрочный интервал
            spacing.set(qn("w:lineRule"), "exact")
            p_pr.append(spacing)
            
            # Убираем отступы
            existing_ind = p_pr.find(qn("w:ind"))
            if existing_ind is not None:
                p_pr.remove(existing_ind)
            ind = OxmlElement("w:ind")
            ind.set(qn("w:left"), "0")
            ind.set(qn("w:firstLine"), "0")
            p_pr.append(ind)
            
            previous_blank = False
            previous_was_parameter_header = True
            index += 1
            continue

        # Проверяем и переводим английские заголовки в обычных строках
        translated_line = line
        # Проверяем на наличие английских заголовков (с двоеточием)
        for eng_base, rus_base in [("Parameters", "Параметры"), ("Returns", "Возвращает"), ("Raises", "Вызывает")]:
            # Ищем паттерн - заголовок с двоеточием
            pattern = re.compile(rf'\b{re.escape(eng_base)}\s*:', re.IGNORECASE)
            if pattern.search(stripped):
                # Проверяем, что русский вариант еще не присутствует
                if rus_base not in stripped:
                    translated_line = pattern.sub(f"{rus_base}:", line)
                    break

        document.add_paragraph(translated_line)
        previous_blank = False
        previous_was_parameter_header = False
        index += 1

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

def add_table_from_markdown(document: Document, table_lines: List[str], no_borders: bool = False) -> None:
    """
    Добавить в DOCX таблицу, созданную из Markdown-синтаксиса.

    Args:
        document: Экземпляр docx.Document.
        table_lines: Набор строк таблицы (включая заголовок).
        no_borders: Если True, таблица создается без рамок.
    """
    parsed_rows: List[List[str]] = []
    for idx, raw_line in enumerate(table_lines):
        cells = [cell.strip() for cell in raw_line.strip().strip("|").split("|")]
        if idx == 1 and all(set(cell) <= {"-", ":", " "} for cell in cells):
            continue  # markdown separator row
        parsed_rows.append(cells)

    if not parsed_rows:
        return

    column_count = len(parsed_rows[0])
    table = document.add_table(rows=len(parsed_rows), cols=column_count)
    try:
        table.style = "Table Grid"
    except ValueError:
        # Оставляем стиль по умолчанию, если Table Grid недоступен
        pass
    
    if no_borders:
        remove_table_borders(table)
    else:
        apply_table_borders(table)
    
    for row_index, row_values in enumerate(parsed_rows):
        for col_index in range(column_count):
            value = row_values[col_index] if col_index < len(row_values) else ""
            table.rows[row_index].cells[col_index].text = value

def ensure_code_style(document: Document) -> None:
    """
    Создать стиль кода, если его еще нет в документе.
    
    Args:
        document: Экземпляр docx.Document.
    """
    try:
        document.styles["Code"]
    except KeyError:
        code_style = document.styles.add_style("Code", WD_STYLE_TYPE.PARAGRAPH)
        code_style.font.name = "Courier New"
        code_style.font.size = Pt(10)
        code_style.paragraph_format.space_before = Pt(6)
        code_style.paragraph_format.space_after = Pt(6)
        code_style.paragraph_format.left_indent = Pt(12)
        code_style.paragraph_format.right_indent = Pt(12)
        
        shading_elm = OxmlElement("w:shd")
        shading_elm.set(qn("w:fill"), "F5F5F5")
        shading_elm.set(qn("w:val"), "clear")
        code_style.element.get_or_add_pPr().append(shading_elm)

def add_list_item(document: Document, line: str, stripped: str, reduce_indent: bool = False) -> None:
    """
    Добавить пункт списка с уменьшенными интервалами.

    Args:
        document: Экземпляр docx.Document.
        line: Полная строка с маркером списка.
        stripped: Строка без начальных пробелов.
        reduce_indent: Если True, уменьшить отступ (для элементов после Параметры/Возвращает/Вызывает).
    """
    text = stripped[2:].strip()
    
    leading_spaces = len(line) - len(line.lstrip())
    is_nested = leading_spaces >= 2
    
    paragraph = document.add_paragraph()
    p_pr = paragraph._element.get_or_add_pPr()
    
    # Используем одинаковое форматирование для всех элементов списка
    # Всегда используем нумерацию Word для единообразия
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0" if not is_nested else "1")
    num_pr.append(ilvl)
    
    num_id = OxmlElement("w:numId")
    num_id.set(qn("w:val"), "1")
    num_pr.append(num_id)
    p_pr.append(num_pr)
    
    existing_ind = p_pr.find(qn("w:ind"))
    if existing_ind is not None:
        p_pr.remove(existing_ind)
    
    ind = OxmlElement("w:ind")
    if reduce_indent:
        # Для элементов после Параметры/Возвращает/Вызывает устанавливаем одинаковый отступ для всех элементов
        ind.set(qn("w:left"), "0")
        ind.set(qn("w:hanging"), "6")  # Минимальный отступ для маркера
        ind.set(qn("w:firstLine"), "-6")  # Отрицательный отступ для первой строки (маркер) - одинаково для всех
    elif is_nested:
        ind.set(qn("w:left"), "72")
        ind.set(qn("w:hanging"), "6")  # Минимальный отступ от маркера
        ind.set(qn("w:firstLine"), "-6")  # Отрицательный отступ для первой строки
    else:
        ind.set(qn("w:left"), "0")
        ind.set(qn("w:hanging"), "6")  # Минимальный отступ от маркера
        ind.set(qn("w:firstLine"), "-6")  # Отрицательный отступ для первой строки
    p_pr.append(ind)
    
    existing_spacing = p_pr.find(qn("w:spacing"))
    if existing_spacing is not None:
        p_pr.remove(existing_spacing)
    
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "0")  # Убираем отступ после всех элементов
    spacing.set(qn("w:line"), "240")
    spacing.set(qn("w:lineRule"), "exact")
    p_pr.append(spacing)
    
    # Всегда используем текст без ручного маркера - Word добавит маркер через нумерацию
    run = paragraph.add_run(text)

def add_code_block(document: Document, code_text: str) -> None:
    """
    Добавить в DOCX блок кода с моноширинным шрифтом.

    Args:
        document: Экземпляр docx.Document.
        code_text: Текст кода.
    """
    paragraph = document.add_paragraph()
    
    p_pr = paragraph._element.get_or_add_pPr()
    
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "F5F5F5")
    shading.set(qn("w:val"), "clear")
    p_pr.append(shading)
    
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "120")
    spacing.set(qn("w:after"), "120")
    p_pr.append(spacing)
    
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "240")
    ind.set(qn("w:right"), "240")
    p_pr.append(ind)
    
    run = paragraph.add_run(code_text)
    r_pr = run._element.get_or_add_rPr()
    
    r_font = OxmlElement("w:rFonts")
    r_font.set(qn("w:ascii"), "Courier New")
    r_font.set(qn("w:hAnsi"), "Courier New")
    r_font.set(qn("w:cs"), "Courier New")
    r_pr.append(r_font)
    
    r_size = OxmlElement("w:sz")
    r_size.set(qn("w:val"), "20")
    r_pr.append(r_size)
    
    r_size_cs = OxmlElement("w:szCs")
    r_size_cs.set(qn("w:val"), "20")
    r_pr.append(r_size_cs)

def remove_table_borders(table: Table) -> None:
    """
    Убрать все границы таблицы.

    Args:
        table: Таблица python-docx.
    """
    tbl = table._tbl
    tbl_props = tbl.tblPr
    if tbl_props is None:
        tbl_props = OxmlElement("w:tblPr")
        tbl.append(tbl_props)

    existing = tbl_props.find(qn("w:tblBorders"))
    if existing is not None:
        tbl_props.remove(existing)
    
    # Создаем пустые границы (без линий)
    borders = OxmlElement("w:tblBorders")
    for border_name in ("top", "bottom", "left", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{border_name}")
        element.set(qn("w:val"), "nil")
        element.set(qn("w:sz"), "0")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "auto")
        borders.append(element)
    tbl_props.append(borders)

def apply_table_borders(table: Table) -> None:
    """
    Насильно прорисовать границы таблицы, чтобы Word отображал рамки.

    Args:
        table: Таблица python-docx.
    """
    tbl = table._tbl
    tbl_props = tbl.tblPr
    if tbl_props is None:
        tbl_props = OxmlElement("w:tblPr")
        tbl.append(tbl_props)

    existing = tbl_props.find(qn("w:tblBorders"))
    if existing is not None:
        tbl_props.remove(existing)

    borders = OxmlElement("w:tblBorders")
    for border_name in ("top", "bottom", "left", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{border_name}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "8")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")
        borders.append(element)

    tbl_props.append(borders)

